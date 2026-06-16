"""Convert analysis Markdown into a formatted Word document."""

from __future__ import annotations

import io
import re
from datetime import date

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


def _set_font(style, name: str, size: int, bold: bool = False, color: str | None = None) -> None:
    style.font.name = name
    style.font.size = Pt(size)
    style.font.bold = bold
    rpr = style.element.get_or_add_rPr()
    rpr.get_or_add_rFonts().set(qn("w:eastAsia"), name)
    if color:
        style.font.color.rgb = RGBColor.from_string(color)


def create_document() -> Document:
    """Create a styled document with a cover page."""
    doc = Document()
    section = doc.sections[0]
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1.25)
    section.right_margin = Inches(1.25)

    normal = doc.styles["Normal"]
    _set_font(normal, "Yu Mincho", 11)
    normal.paragraph_format.line_spacing = 1.5
    normal.paragraph_format.space_after = Pt(6)
    for name, size, color, after in (
        ("Heading 1", 16, "1F3864", 12),
        ("Heading 2", 14, "2F5496", 8),
        ("Heading 3", 12, "1F3864", 6),
    ):
        style = doc.styles[name]
        _set_font(style, "Yu Mincho", size, True, color)
        style.paragraph_format.space_after = Pt(after)

    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title.paragraph_format.space_before = Pt(180)
    run = title.add_run("日本語テキスト分析レポート")
    run.bold = True
    run.font.size = Pt(22)
    run.font.name = "Yu Mincho"
    run._element.get_or_add_rPr().get_or_add_rFonts().set(qn("w:eastAsia"), "Yu Mincho")
    created = doc.add_paragraph(date.today().isoformat())
    created.alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_page_break()
    return doc


def _shade(cell, fill: str) -> None:
    properties = cell._tc.get_or_add_tcPr()
    shading = properties.find(qn("w:shd"))
    if shading is None:
        shading = OxmlElement("w:shd")
        properties.append(shading)
    shading.set(qn("w:fill"), fill)


def _set_borders(table) -> None:
    properties = table._tbl.tblPr
    borders = properties.find(qn("w:tblBorders"))
    if borders is None:
        borders = OxmlElement("w:tblBorders")
        properties.append(borders)
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        tag = borders.find(qn(f"w:{edge}"))
        if tag is None:
            tag = OxmlElement(f"w:{edge}")
            borders.append(tag)
        tag.set(qn("w:val"), "single")
        tag.set(qn("w:sz"), "4")
        tag.set(qn("w:color"), "9DC3E6")


def _table_values(line: str) -> list[str]:
    return [cell.strip() for cell in line.strip().strip("|").split("|")]


def add_markdown_table(doc: Document, table_lines: list[str]) -> None:
    """Add a styled Word table from consecutive Markdown table lines."""
    if len(table_lines) < 2:
        return
    rows = [_table_values(line) for line in table_lines]
    if len(rows) > 1 and all(re.fullmatch(r":?-{3,}:?", cell.replace(" ", "")) for cell in rows[1]):
        rows.pop(1)
    column_count = max(len(row) for row in rows)
    table = doc.add_table(rows=len(rows), cols=column_count)
    table.autofit = True
    _set_borders(table)
    for row_index, values in enumerate(rows):
        for column_index in range(column_count):
            cell = table.cell(row_index, column_index)
            cell.text = values[column_index] if column_index < len(values) else ""
            _shade(cell, "BDD7EE" if row_index == 0 else ("F2F2F2" if row_index % 2 == 0 else "FFFFFF"))
            for run in cell.paragraphs[0].runs:
                run.bold = row_index == 0
                run.font.name = "Yu Mincho"
                run._element.get_or_add_rPr().get_or_add_rFonts().set(qn("w:eastAsia"), "Yu Mincho")


def _add_inline(paragraph, text: str) -> None:
    """Add text while applying **bold** and `code` inline formatting."""
    position = 0
    for match in re.finditer(r"(\*\*(.+?)\*\*|`(.+?)`)", text):
        if match.start() > position:
            paragraph.add_run(text[position : match.start()])
        if match.group(2) is not None:
            paragraph.add_run(match.group(2)).bold = True
        else:
            run = paragraph.add_run(match.group(3))
            run.font.name = "Courier New"
        position = match.end()
    if position < len(text):
        paragraph.add_run(text[position:])


def parse_and_add_markdown(doc: Document, markdown_text: str) -> None:
    """Parse supported Markdown constructs and add them to a document."""
    lines = markdown_text.splitlines()
    index = 0
    while index < len(lines):
        line = lines[index].rstrip()
        if line.strip().startswith("|"):
            table_lines = []
            while index < len(lines) and lines[index].strip().startswith("|"):
                table_lines.append(lines[index])
                index += 1
            add_markdown_table(doc, table_lines)
            continue
        stripped = line.strip()
        if not stripped:
            index += 1
            continue
        heading = re.match(r"^(#{1,3})\s+(.+)$", stripped)
        numbered = re.match(r"^\d+\.\s+(.+)$", stripped)
        if heading:
            paragraph = doc.add_heading(level=len(heading.group(1)))
            _add_inline(paragraph, heading.group(2))
        elif stripped.startswith(("- ", "* ")):
            paragraph = doc.add_paragraph(style="List Bullet")
            _add_inline(paragraph, stripped[2:])
        elif numbered:
            paragraph = doc.add_paragraph(style="List Number")
            _add_inline(paragraph, numbered.group(1))
        else:
            paragraph = doc.add_paragraph(style="Normal")
            _add_inline(paragraph, stripped)
        index += 1


def _add_footer(doc: Document) -> None:
    footer = doc.sections[0].footer.paragraphs[0]
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    footer.add_run("English OCR Analyzer | Trang ")
    field = OxmlElement("w:fldSimple")
    field.set(qn("w:instr"), "PAGE")
    footer._p.append(field)


def export_to_docx(markdown_content: str, output_filename: str = "analysis.docx") -> bytes:
    """Export Markdown content to DOCX bytes."""
    if not markdown_content or not markdown_content.strip():
        raise ValueError("Nội dung Markdown không được rỗng.")
    if not output_filename.lower().endswith(".docx"):
        raise ValueError("Tên file xuất phải có phần mở rộng .docx.")
    doc = create_document()
    parse_and_add_markdown(doc, markdown_content)
    _add_footer(doc)
    buffer = io.BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer.getvalue()
