"""Export dialogue practice results to TXT, JSON, and Word (.docx) formats."""

from __future__ import annotations

import io
import json
import copy
from typing import Any

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


def export_dialogue_to_text(result: dict[str, Any]) -> str:
    """Export dialogue as clean plain text."""
    lines = []
    lines.append(f"=== HỘI THOẠI: {result.get('topic', 'Chủ đề')} ===")
    lines.append(f"Ngôn ngữ: {result.get('language', 'N/A')} | Cấp độ: {result.get('level', 'N/A')}")
    lines.append(f"Tình huống: {result.get('situation', 'Thông thường')} | Kính ngữ: {result.get('politeness_level', 'Lịch sự')}")
    if result.get("scenario_description"):
        lines.append(f"Miêu tả hoàn cảnh: {result['scenario_description']}")
    lines.append("-" * 50)
    lines.append("")

    for turn in result.get("dialogue", []):
        lines.append(f"[{turn['speaker']}]: {turn['text']}")
        if turn.get("text_hira"):
            lines.append(f"     Cách đọc: {turn['text_hira']}")
        if turn.get("text_vi"):
            lines.append(f"     Dịch Việt: {turn['text_vi']}")
        if turn.get("highlights"):
            lines.append(f"     Từ mục tiêu: {', '.join(turn['highlights'])}")
        if turn.get("speech_intent"):
            lines.append(f"     Mục đích nói: {turn['speech_intent']}")
        if turn.get("alternative_expression"):
            lines.append(f"     Cách nói khác: {turn['alternative_expression']}")
        lines.append("")

    if result.get("learning_targets"):
        lines.append("=" * 50)
        lines.append("MỤC TIÊU HỌC:")
        for row in result["learning_targets"]:
            if not isinstance(row, dict):
                continue
            lines.append(f"- [{row.get('type', 'vocabulary')}] {row.get('term', '')}")
            if row.get("realized_form"):
                lines.append(f"  Dạng dùng: {row['realized_form']}")
            if row.get("explanation_vi"):
                lines.append(f"  Giải thích: {row['explanation_vi']}")
        lines.append("")

    if result.get("summary"):
        lines.append("=" * 50)
        lines.append("TÓM TẮT TỪ VỰNG & NGỮ PHÁP:")
        lines.append(result["summary"])
        lines.append("")

    if result.get("notes"):
        lines.append("GIẢI THÍCH NGỮ CẢNH:")
        lines.append(result["notes"])

    return "\n".join(lines)


def export_dialogue_to_json(result: dict[str, Any]) -> str:
    """Export the complete portable learning record, excluding only raw model text."""
    clean_dict = copy.deepcopy(result)
    clean_dict.pop("raw_text", None)
    return json.dumps(clean_dict, ensure_ascii=False, indent=2, default=str)


def _shade_cell(cell: Any, color_hex: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), color_hex)
    tc_pr.append(shd)


def export_dialogue_to_docx(result: dict[str, Any]) -> bytes:
    """Export dialogue result to a beautifully formatted Word (.docx) document."""
    doc = Document()

    # Page setup
    for sec in doc.sections:
        sec.top_margin = Inches(0.8)
        sec.bottom_margin = Inches(0.8)
        sec.left_margin = Inches(0.8)
        sec.right_margin = Inches(0.8)

    # Title
    title_p = doc.add_paragraph()
    title_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title_p.add_run(f"📖 BÀI LUYỆN HỘI THOẠI: {result.get('topic', 'Chủ đề')}")
    run.bold = True
    run.font.size = Pt(18)
    run.font.name = "Arial"
    run.font.color.rgb = RGBColor(0, 51, 102)

    # Subtitle / Meta
    meta_p = doc.add_paragraph()
    meta_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    meta_text = (
        f"Ngôn ngữ: {result.get('language', 'Tiếng Nhật')} | Cấp độ: {result.get('level', 'Trung bình')} | "
        f"Tình huống: {result.get('situation', 'Thông thường')} | Kính ngữ: {result.get('politeness_level', 'Lịch sự')}"
    )
    if result.get("scenario_description"):
        meta_text += f"\nMiêu tả hoàn cảnh: {result['scenario_description']}"
    meta_run = meta_p.add_run(meta_text)
    meta_run.font.size = Pt(10)
    meta_run.font.italic = True
    meta_run.font.color.rgb = RGBColor(100, 100, 100)

    doc.add_paragraph().paragraph_format.space_after = Pt(6)

    # Dialogue Table
    table = doc.add_table(rows=1, cols=3)
    table.style = "Table Grid"

    hdr_cells = table.rows[0].cells
    is_english = "english" in str(result.get("language", "")).lower() or "tiếng anh" in str(result.get("language", "")).lower()
    headers = ["Người nói", "Câu tiếng Anh" if is_english else "Câu tiếng Nhật / Cách đọc", "Bản dịch tiếng Việt"]
    widths = [Inches(1.0), Inches(3.2), Inches(2.6)]

    for idx, (hdr_text, w) in enumerate(zip(headers, widths)):
        hdr_cells[idx].width = w
        p = hdr_cells[idx].paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = p.add_run(hdr_text)
        r.bold = True
        r.font.size = Pt(10.5)
        r.font.color.rgb = RGBColor(255, 255, 255)
        _shade_cell(hdr_cells[idx], "1F4E79")

    for i, turn in enumerate(result.get("dialogue", [])):
        row_cells = table.add_row().cells
        for idx, w in enumerate(widths):
            row_cells[idx].width = w

        bg_color = "F2F7FA" if i % 2 == 1 else "FFFFFF"
        for c in row_cells:
            _shade_cell(c, bg_color)

        # Speaker
        spk_p = row_cells[0].paragraphs[0]
        spk_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        role_names = {
            str(row.get("id", "")).upper(): str(row.get("name", "")).strip()
            for row in result.get("roles", []) if isinstance(row, dict)
        }
        speaker_name = role_names.get(str(turn.get("speaker", "")).upper()) or f"Nhân vật {turn.get('speaker', '')}"
        spk_run = spk_p.add_run(speaker_name)
        spk_run.bold = True
        spk_run.font.size = Pt(10)
        spk_run.font.color.rgb = RGBColor(0, 102, 153) if turn["speaker"] == "A" else RGBColor(153, 51, 0)

        # Target / Hira
        jp_p = row_cells[1].paragraphs[0]
        jp_run = jp_p.add_run(turn["text"])
        jp_run.font.size = Pt(11)
        jp_run.bold = True

        if turn.get("text_hira"):
            hira_p = row_cells[1].add_paragraph()
            hira_run = hira_p.add_run(f"({turn['text_hira']})")
            hira_run.font.size = Pt(9.5)
            hira_run.font.italic = True
            hira_run.font.color.rgb = RGBColor(90, 90, 90)

        # Vietnamese
        vi_p = row_cells[2].paragraphs[0]
        vi_run = vi_p.add_run(turn.get("text_vi", ""))
        vi_run.font.size = Pt(10)

    doc.add_paragraph().paragraph_format.space_after = Pt(12)

    # Summary Section
    if result.get("learning_targets"):
        h2 = doc.add_heading("Mục tiêu học", level=2)
        h2.style.font.name = "Arial"
        h2.style.font.color.rgb = RGBColor(0, 51, 102)
        for target in result["learning_targets"]:
            if not isinstance(target, dict):
                continue
            p = doc.add_paragraph(style="List Bullet")
            p.add_run(f"[{target.get('type', 'vocabulary')}] {target.get('term', '')}").bold = True
            detail = " · ".join(
                value for value in (target.get("realized_form"), target.get("explanation_vi")) if value
            )
            if detail:
                p.add_run("\n" + detail)

    if result.get("summary"):
        h2 = doc.add_heading("📚 Tóm tắt Từ vựng & Ngữ pháp", level=2)
        h2.style.font.name = "Arial"
        h2.style.font.color.rgb = RGBColor(0, 51, 102)

        sum_p = doc.add_paragraph()
        sum_p.add_run(result["summary"]).font.size = Pt(10)

    # Notes Section
    if result.get("notes"):
        h2 = doc.add_heading("💡 Ghi chú ngữ cảnh", level=2)
        h2.style.font.name = "Arial"
        h2.style.font.color.rgb = RGBColor(0, 51, 102)

        notes_p = doc.add_paragraph()
        notes_p.add_run(result["notes"]).font.size = Pt(10)

    buffer = io.BytesIO()
    doc.save(buffer)
    return buffer.getvalue()
