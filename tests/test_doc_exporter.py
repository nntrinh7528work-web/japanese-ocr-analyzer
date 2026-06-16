import io

from docx import Document

from modules.doc_exporter import export_to_docx


SAMPLE = """# PHÂN TÍCH VĂN BẢN TIẾNG NHẬT

## 1. XÁC NHẬN VĂN BẢN GỐC
日本は島国です。

**Tóm tắt:** Bài viết nói về `日本`.

1. Mục thứ nhất
- Ghi chú

| # | Từ gốc | Phiên âm | Nghĩa |
|---|---|---|---|
| 1 | 日本 | にほん | Nhật Bản |
| 2 | 島国 | しまぐに | Đảo quốc |
"""


def test_export_docx_is_valid_and_formatted(tmp_path):
    docx_bytes = export_to_docx(SAMPLE, "test_output.docx")
    output = tmp_path / "test_output.docx"
    output.write_bytes(docx_bytes)
    doc = Document(io.BytesIO(docx_bytes))

    assert len(docx_bytes) > 1000
    assert len(doc.tables) == 1
    assert doc.tables[0].rows[0].cells[1].text == "Từ gốc"
    assert any(paragraph.style.name == "Heading 1" for paragraph in doc.paragraphs)
    assert "Japanese / English OCR Analyzer" in doc.sections[0].footer.paragraphs[0].text
